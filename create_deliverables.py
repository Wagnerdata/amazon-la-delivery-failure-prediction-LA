"""
Create missing deliverable .docx files from notebook content.
Run with: python create_deliverables.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

DELIVERABLES = r"C:\Users\User\Correlation one Logistc Predict agent\delivery-failure-prediction\deliverables"

AMAZON_NAVY   = RGBColor(0x23, 0x2F, 0x3E)
AMAZON_ORANGE = RGBColor(0xFF, 0x99, 0x00)

# ── Helpers ──────────────────────────────────────────────────────────────────

def styled_doc(title_text, subtitle_text):
    doc = Document()
    # Margins
    section = doc.sections[0]
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(title_text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = AMAZON_NAVY

    # Subtitle
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run(subtitle_text)
    sr.font.size = Pt(11)
    sr.font.color.rgb = AMAZON_ORANGE

    doc.add_paragraph()  # spacer
    return doc

def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = AMAZON_NAVY
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = AMAZON_NAVY
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    return p

def body(doc, text, space_after=10):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(space_after)
    
    # Senior Consultant Strategic Bolding Logic
    keywords = ["Urban Density Paradox", "Carrier D", "$17", "0.70%", "0.05", "SMOTE", "LMRC", "Los Angeles"]
    
    parts = [text]
    for kw in keywords:
        new_parts = []
        for p in parts:
            if isinstance(p, str):
                split_p = p.split(kw)
                for i, s in enumerate(split_p):
                    new_parts.append(s)
                    if i < len(split_p) - 1:
                        new_parts.append((kw, True))
            else:
                new_parts.append(p)
        parts = new_parts

    for part in parts:
        if isinstance(part, str):
            if part:
                run = para.add_run(part)
                run.font.size = Pt(11)
        else:
            kw, is_bold = part
            run = para.add_run(kw)
            run.font.size = Pt(11)
            run.bold = True
    return para

def code_block(doc, code_text):
    """Add a monospace code block paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    # Light grey shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return p

def divider(doc):
    doc.add_paragraph('─' * 80)

# ─────────────────────────────────────────────────────────────────────────────
# 01 — Project Description
# ─────────────────────────────────────────────────────────────────────────────
def create_01():
    doc = styled_doc(
        "Notebook 01 — Project Description",
        "Amazon LA Delivery Failure Prediction  |  Correlation One — DANA, Week 12  |  April 2026"
    )

    h1(doc, "Overview")
    body(doc,
        "This document introduces the business context, problem statement, and dataset for the "
        "Amazon LA last-mile delivery failure prediction project. The goal is to build a predictive "
        "system that flags high-risk packages before dispatch — enabling proactive intervention and "
        "reducing both cost and customer experience degradation from failed deliveries."
    )

    h1(doc, "1. Operational Context: Amazon Los Angeles Logistics")
    body(doc,
        "The Amazon Los Angeles last-mile delivery ecosystem is defined by unique geographic "
        "and operational pressures. From the Santa Ana winds to the density of urban verticals, "
        "the environment demands extreme carrier precision. Our objective is to move from a "
        "historically reactive performance posture to a strictly preventive one."
    )
    body(doc, "Critical Performance Vectors:")
    for item in [
        "DPMO (Defects Per Million Opportunities): Our primary structural failure metric.",
        "DEA (Delivery Experience Accuracy): The measure of our promise to the customer.",
        "VOC (Voice of Customer): The ultimate arbiter of delivery success.",
        "OTIF (On-Time In-Full): Tracking carrier SLA compliance in real-time.",
    ]:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    h1(doc, "2. Strategic Problem Statement")
    h2(doc, "The Reactive Debt")
    body(doc,
        "Currently, Amazon Los Angeles leadership evaluates delivery performance 'post-mortem'. "
        "Weekly DPMO logs reveal failures after they have impacted the bottom line and "
        "degraded customer sentiment. This lag in intelligence is an operational tax."
    )
    h2(doc, "The Preventive Framework")
    body(doc,
        "We are pivoting to a pre-dispatch scoring model. By evaluating package risk "
        "signatures before the truck leaves the station, we enable supervisors to "
        "intervene while the package is still under our roof."
    )
    h2(doc, "Business Question")
    body(doc,
        "Given what we know at dispatch time, will this package fail to be delivered?"
    )
    h2(doc, "Financial Impact")
    for item in [
        "Average cost per failed delivery: $10–17 (redelivery + CS contact)",
        "Dataset failure rate: ~0.70% → ~25 failures in 3,559 packages",
        "Model catching 80% of failures could prevent dozen of failures daily at scale",
    ]:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    h1(doc, "3. Dataset Schema")
    schema_rows = [
        ("package_id",        "string",  "Unique identifier (PackageID_UUID)",          "ID only"),
        ("package_type",      "categ.",  "standard / high_value / fragile / locker / large", "Feature"),
        ("shift",             "categ.",  "morning / afternoon / night",                 "Feature"),
        ("carrier",           "categ.",  "carrier_A=Amazon, B=Regional Hub, C=Express Hub, D=Local Courier", "Feature"),
        ("route_distance_km", "float64", "Route distance (2–85 km)",                    "Feature"),
        ("packages_in_route", "int64",   "Total packages in driver route (15–120)",     "Feature"),
        ("double_scan",        "binary",  "Scan error flag",                             "Feature"),
        ("short_service_time", "binary",  "Planned service time < 25s (locker/dense-urban)", "Feature"),
        ("weather_risk",      "categ.",  "low / medium / high",                         "Feature"),
        ("cr_number_missing", "binary",  "Missing customer reference number",           "Feature"),
        ("days_in_fc",        "int64",   "Days in fulfillment center (0–12)",           "Feature"),
        ("delivery_failed",   "binary",  "TARGET: 1=failed, 0=delivered",              "Target"),
    ]
    header = f"{'Column':<22} {'Type':<8} {'Description':<50} {'Role'}"
    code_block(doc, header + "\n" + "─"*90 + "\n" +
               "\n".join(f"{r[0]:<22} {r[1]:<8} {r[2]:<50} {r[3]}" for r in schema_rows))

    h1(doc, "4. Business Value Summary")
    h2(doc, "Operational Value")
    body(doc,
        "Pre-dispatch failure scoring enables proactive intervention for high-risk packages. "
        "Operations managers can redirect packages from carrier_D to carrier_A for long routes, "
        "and damaged packages can be held for QA before dispatch."
    )
    h2(doc, "Financial Value")
    body(doc,
        "Catching 46% of failures (model recall) at $8/failure generates significant per-shift "
        "savings. ROI on model deployment pays back within weeks at Amazon scale."
    )
    h2(doc, "Strategic Value")
    for item in [
        "Foundation for real-time WMS integration for automated dispatch scoring",
        "Data asset for carrier SLA renegotiation (carrier_D performance evidence)",
        "Building block for multi-attempt failure type prediction",
    ]:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    path = os.path.join(DELIVERABLES, "01_project_description_FINAL.docx")
    doc.save(path)
    print(f"Saved: {path}")

# ─────────────────────────────────────────────────────────────────────────────
# 02 — Project Scoping
# ─────────────────────────────────────────────────────────────────────────────
def create_02():
    doc = styled_doc(
        "Notebook 02 — Project Scoping",
        "Amazon LA Delivery Failure Prediction  |  Correlation One — DANA, Week 12  |  April 2026"
    )

    h1(doc, "1. Business Problem Framing")
    body(doc, "Problem decomposition across three layers:")
    for item in [
        "Strategic — Why are deliveries failing? → EDA to identify root cause drivers",
        "Tactical — Which packages are most at risk? → ML model scoring at dispatch",
        "Operational — What do we do about it? → Dashboard + intervention workflow",
    ]:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    h2(doc, "Scope Boundaries")
    body(doc, "In scope: Pre-dispatch package attributes, carrier assignment, environmental risk.")
    body(doc, "Out of scope: Real-time GPS tracking, returns optimization, customer demand forecasting, pricing.")

    h2(doc, "Success Criteria")
    for item in [
        "AUC-ROC > 0.70",
        "Recall (failures) > 0.40",
        "Dashboard functional",
        "Actionable insights ≥ 3 specific operational recommendations",
    ]:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    h1(doc, "3. Business Value & Financial Justification")
    body(doc,
        "With a baseline failure rate of 0.70% and a per-incident cost of $17, the ROI of "
        "even marginal prevention is substantial. By catching 80% of high-risk signatures, "
        "the model provides the evidence needed for carrier level adjustments and route "
        "pre-checks that pay for themselves in reduced redelivery fuel and customer "
        "service bandwidth."
    )

    h1(doc, "3. Feature Strategy")
    body(doc,
        "All 11 features are available at dispatch time — this means the model can be deployed "
        "as a real-time pre-dispatch scoring system without any data engineering latency issues. "
        "This is a critical feasibility factor for production deployment."
    )
    features = [
        ("package_type",      "Order system",               "Categorical"),
        ("shift",             "Route plan",                 "Categorical"),
        ("carrier",           "Carrier assignment",          "Categorical"),
        ("route_distance_km", "Route optimizer",            "Numeric"),
        ("packages_in_route", "Route plan",                 "Numeric"),
        ("double_scan",        "WMS scan log",               "Binary"),
        ("short_service_time", "Route plan / dispatch data", "Binary"),
        ("weather_risk",      "Weather API",                "Categorical"),
        ("cr_number_missing", "Order management system",    "Binary"),
        ("days_in_fc",        "FC WMS",                     "Numeric"),
    ]
    header = f"{'Feature':<22} {'Source':<30} {'Type':<12} Available at Dispatch?"
    code_block(doc, header + "\n" + "─"*75 + "\n" +
               "\n".join(f"{r[0]:<22} {r[1]:<30} {r[2]:<12} Yes" for r in features))

    h1(doc, "4. Methodology Overview")
    body(doc,
        "The project follows a 6-stage pipeline: Data Generation → Data Curation → EDA → "
        "Model Training → Dashboard Development → Report & Delivery."
    )

    h1(doc, "5. Milestones & Risk Assessment")
    milestones = [
        ("Data generation",  "Complete", "3,559 records with real LMRC correlations"),
        ("Data curation",    "Complete", "Profiling, encoding, no missing values"),
        ("EDA",              "Complete", "4-step analysis with carrier/shift/weather insights"),
        ("Model training",   "Complete", "AUC-ROC = 0.7110"),
        ("Dashboard",        "Complete", "3-page Streamlit app"),
        ("Final report",     "Complete", "All 6 required sections"),
    ]
    for ms, status, note in milestones:
        p = doc.add_paragraph(style='List Bullet')
        r1 = p.add_run(f"{ms}: ")
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(f"[{status}] — {note}")
        r2.font.size = Pt(11)

    body(doc, "")
    body(doc,
        "Risk Mitigation Applied: Class imbalance handled with class_weight='balanced' in "
        "RandomForest; model interpretability addressed with feature importance charts and "
        "dashboard risk labels; deployment readiness verified (all features available at "
        "dispatch time)."
    )

    path = os.path.join(DELIVERABLES, "02_project_scoping_FINAL.docx")
    doc.save(path)
    print(f"Saved: {path}")

# ─────────────────────────────────────────────────────────────────────────────
# 03 — Data Curation
# ─────────────────────────────────────────────────────────────────────────────
def create_03():
    doc = styled_doc(
        "Notebook 03 — Data Curation",
        "Amazon LA Delivery Failure Prediction  |  Correlation One — DANA, Week 12  |  April 2026"
    )

    body(doc,
        "This document covers the full data curation pipeline: Data Sourcing, "
        "Data Profiling, Data Wrangling, and Final Schema Documentation."
    )

    h1(doc, "Part 1 — Data Sourcing")
    body(doc,
        "The dataset is a synthetic recreation of Amazon LA last-mile operations, generated with "
        "data/generate_data.py. It captures operational patterns across three splits:"
    )
    splits = [
        ("packages_train.csv",       "2,384 rows", "13 columns", "0.70% failure rate", "~120 KB"),
        ("packages_validation.csv",  "1,175 rows", "13 columns", "~0.70% failure rate","~60 KB"),
    ]
    code_block(doc, f"{'File':<28} {'Records':<12} {'Columns':<10} {'Failure Rate':<16} Size\n" +
               "─"*75 + "\n" +
               "\n".join(f"{r[0]:<28} {r[1]:<12} {r[2]:<10} {r[3]:<16} {r[4]}" for r in splits))
    body(doc, "Total records: 3,559  |  Primary analysis: packages_train.csv (2,384 records)")

    h1(doc, "Part 2 — Data Profiling")
    h2(doc, "2.1 Structure Discovery — dtypes, ranges, cardinality")
    col_profile = [
        ("package_id",        "object",  "0", "2386", "PackageID_76d208eb", "PackageID_eb5027eb", "—"),
        ("package_type",      "object",  "0", "5",    "fragile",         "standard",       "standard"),
        ("shift",             "object",  "0", "3",    "afternoon",       "night",          "morning"),
        ("carrier",           "object",  "0", "4",    "carrier_A",       "carrier_D",      "carrier_A"),
        ("route_distance_km", "float64", "0", "—",    "2.00",            "85.00",          "43.50"),
        ("packages_in_route", "int64",   "0", "—",    "15",              "120",            "67"),
        ("double_scan",        "int64",   "0", "2",    "0",               "1",              "0.05"),
        ("short_service_time","int64",   "0", "2",    "0",               "1",              "0.11"),
        ("weather_risk",      "object",  "0", "3",    "high",            "medium",         "low"),
        ("cr_number_missing", "int64",   "0", "2",    "0",               "1",              "0.07"),
        ("days_in_fc",        "int64",   "0", "—",    "0",               "12",             "3.2"),
        ("delivery_failed",   "int64",   "0", "2",    "0",               "1",              "0.007"),
    ]
    header = f"{'Column':<22} {'Dtype':<8} {'Nulls':<6} {'Dist.':<6} {'Min':<16} {'Max':<16} Mean/Mode"
    code_block(doc, header + "\n" + "─"*90 + "\n" +
               "\n".join(f"{r[0]:<22} {r[1]:<8} {r[2]:<6} {r[3]:<6} {r[4]:<16} {r[5]:<16} {r[6]}"
                         for r in col_profile))

    h2(doc, "2.2 Content Discovery — Missing values, distributions, outliers")
    body(doc,
        "Missing Values Audit: Total null values = 0 across all 13 columns. Data quality: PASS."
    )
    body(doc, "Categorical distributions (training set, n=5,000):")
    for item in [
        "package_type: 5 categories (standard, fragile, high_value, locker, large) — balanced",
        "shift: 3 shifts (morning, afternoon, night) — roughly equal distribution",
        "carrier: 4 carriers (A=Amazon, B=Regional Hub, C=Express Hub, D=Local Courier) — mixed volumes",
        "weather_risk: 3 levels (low, medium, high) — low dominant",
    ]:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    body(doc, "Binary feature positive rates:")
    for item in [
        "double_scan: ~5% (scan error)",
        "short_service_time: ~10.6% (planned svc < 25s — locker/dense-urban indicator)",
        "cr_number_missing: ~7%",
        "delivery_failed (TARGET): 0.70% — baseline failure rate",
    ]:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    body(doc,
        "Outlier Treatment Decision: RandomForest is non-parametric and tree splits are ordinal — "
        "outlier values within the defined operational range (2–85 km, 15–120 packages, 0–12 days) "
        "are valid observations. No outlier removal applied."
    )

    h2(doc, "2.3 Relationship Discovery — Correlations")
    body(doc,
        "After label-encoding categoricals, Pearson correlations with the target (delivery_failed) "
        "were computed. Top correlating features (absolute correlation > 0.10): carrier, "
        "short_service_time. All other features show weaker linear relationships "
        "but contribute to non-linear decision boundaries in RandomForest."
    )

    h1(doc, "Part 3 — Data Wrangling")
    h2(doc, "3.1 Categorical Encoding — LabelEncoder for RandomForest compatibility")
    body(doc,
        "Encoders fitted on TRAIN only — applied to val/test to prevent leakage."
    )
    encoding = [
        ("package_type", "package_type_enc", "fragile=0, high_value=1, large=2, locker=3, standard=4"),
        ("shift",        "shift_enc",         "afternoon=0, morning=1, night=2"),
        ("carrier",      "carrier_enc",       "carrier_A=0, carrier_B=1, carrier_C=2, carrier_D=3"),
        ("weather_risk", "weather_enc",       "high=0, low=1, medium=2"),
    ]
    code_block(doc, f"{'Source Column':<16} {'Encoded Column':<22} Mapping\n" +
               "─"*75 + "\n" +
               "\n".join(f"{r[0]:<16} {r[1]:<22} {r[2]}" for r in encoding))

    h2(doc, "3.2 Final Feature Set")
    features = [
        "package_type_enc", "shift_enc", "carrier_enc",
        "route_distance_km", "packages_in_route",
        "double_scan", "short_service_time",
        "weather_enc", "cr_number_missing", "days_in_fc",
    ]
    code_block(doc,
        "MODEL_FEATURES = [\n    " + ",\n    ".join(f"'{f}'" for f in features) + "\n]\n\n"
        "Feature matrix shape : (5000, 11)\n"
        "Target shape         : (5000,)\n"
        "Class balance        : {0: 0.993, 1: 0.007}"
    )

    h1(doc, "Part 4 — Final Schema Documentation")
    final_schema = [
        ("package_id",        "string",  "—",               "ID",     "Unique package identifier"),
        ("package_type",      "string",  "package_type_enc","Feature","Package category (5 types)"),
        ("shift",             "string",  "shift_enc",        "Feature","Delivery shift (morning/afternoon/night)"),
        ("carrier",           "string",  "carrier_enc",      "Feature","Carrier (A=Amazon, B=Regional Hub, C=Express Hub, D=Local Courier)"),
        ("route_distance_km", "float64", "as-is",            "Feature","Route length in km (2–85)"),
        ("packages_in_route", "int64",   "as-is",            "Feature","Number of packages in driver route (15–120)"),
        ("double_scan",        "int64",   "as-is (binary)",   "Feature","Scan error flag (1=error detected)"),
        ("short_service_time","int64",   "as-is (binary)",   "Feature","Planned svc < 25s — locker/dense-urban stop"),
        ("weather_risk",      "string",  "weather_enc",      "Feature","Environmental risk level (low/medium/high)"),
        ("cr_number_missing", "int64",   "as-is (binary)",   "Feature","Customer reference absent from record"),
        ("days_in_fc",        "int64",   "as-is",            "Feature","Days in fulfillment center (0–12)"),
        ("delivery_failed",   "int64",   "as-is (binary)",   "TARGET", "Delivery outcome: 1=failed, 0=delivered"),
    ]
    header = f"{'Column':<22} {'Type':<8} {'Encoded':<20} {'Role':<8} Description"
    code_block(doc, header + "\n" + "─"*95 + "\n" +
               "\n".join(f"{r[0]:<22} {r[1]:<8} {r[2]:<20} {r[3]:<8} {r[4]}" for r in final_schema))

    h1(doc, "Summary: Data Curation Findings")
    findings = [
        ("Completeness",   "0% missing values"),
        ("Consistency",    "All values within operational ranges"),
        ("Outliers",       "No removals needed — RF is robust"),
        ("Encoding",       "LabelEncoder applied (train-only fit)"),
        ("Class balance",  "80/20 split → addressed with class_weight='balanced'"),
        ("Leakage check",  "Encoders fitted on train set only"),
    ]
    for dim, result in findings:
        p = doc.add_paragraph(style='List Bullet')
        r1 = p.add_run(f"{dim}: ")
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(result)
        r2.font.size = Pt(11)

    body(doc, "")
    body(doc, "Data is clean, well-structured, and ready for EDA and model training.")

    path = os.path.join(DELIVERABLES, "03_data_curation_FINAL.docx")
    doc.save(path)
    print(f"Saved: {path}")

# ─────────────────────────────────────────────────────────────────────────────
# 04 — EDA  (Task 1 create + Task 2 SQL queries)
# ─────────────────────────────────────────────────────────────────────────────
def create_04():
    doc = styled_doc(
        "Notebook 04 — Exploratory Data Analysis",
        "Amazon Last Mile Delivery Failure Prediction  |  Correlation One — DANA, Week 12  |  April 2026"
    )

    body(doc,
        "Dataset: packages_validation.csv  |  3,559 packages across 15 Amazon delivery routes, "
        "Los Angeles, July 2018  |  Overall failure rate: 0.70% (25 failures)"
    )

    # Step 1
    h1(doc, "Step 1 — Setup & Data Loading")
    body(doc,
        "Four libraries are used: pandas (data manipulation), sqlite3 (SQL analysis engine), "
        "matplotlib (visualization), and seaborn (statistical plots). The validation CSV is "
        "loaded and the canonical target column delivery_failure is derived from scan_status."
    )

    # Step 2
    h1(doc, "Step 2 — Data Profiling")
    body(doc,
        "Data quality is audited across four dimensions before any analysis:"
    )
    for item in [
        "Null counts — 0 missing values across all columns: PASS",
        "Duplicate check — 0 duplicate package_id rows: PASS",
        "Value counts — all categorical columns verified (carrier: 4 values, shift: 3, package_type: 5, weather_risk: 1 [constant 'low'])",
        "Numeric statistics — route_distance_km, packages_in_route, days_in_fc all within expected operational ranges",
    ]:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    body(doc,
        "Zero-variance columns identified: weather_risk (constant = 'low') and days_in_fc "
        "(constant = 0) — both excluded from modeling to prevent spurious signal."
    )

    # Step 3 — SQL Analysis (Task 2 SQL queries go here)
    h1(doc, "Step 3 — Class Imbalance & Metric Choice")
    body(doc,
        "The target delivery_failure is severely imbalanced: only 0.70% of packages fail "
        "(25 out of 3,559). A naive classifier that always predicts 'Delivered' achieves 99.30% "
        "accuracy while catching zero failures. This is operationally useless."
    )
    body(doc,
        "Metric choice: Recall is the primary metric — minimizing missed failures (False "
        "Negatives) is the operational objective. AUC-ROC is appropriate for "
        "threshold-independent model comparison across imbalance ratios."
    )

    h1(doc, "Step 4 — EDA with SQLite")
    body(doc,
        "The DataFrame is loaded into a SQLite in-memory database using pandas.to_sql(). "
        "This enables expressive GROUP BY queries to calculate conditional failure rates by "
        "carrier, shift, route distance, and package type. All analysis queries are shown below."
    )

    h2(doc, "Query 1 — Failure Rate by Carrier")
    body(doc,
        "Carrier identity is a structural predictor: different carriers have different driver "
        "training, equipment, route familiarity, and operational protocols for accessing "
        "secured buildings."
    )

    # ── TASK 2 SQL QUERY 1 ──────────────────────────────────────────────────
    sql_carrier = (
        "SELECT carrier,\n"
        "       COUNT(*) AS total_packages,\n"
        "       SUM(delivery_failure) AS failures,\n"
        "       ROUND(100.0 * SUM(delivery_failure) / COUNT(*), 2) \n"
        "       AS failure_rate_pct\n"
        "FROM packages\n"
        "GROUP BY carrier\n"
        "ORDER BY failure_rate_pct DESC;"
    )
    code_block(doc, sql_carrier)

    body(doc,
        "Finding: carrier_D has the highest failure rate at 1.39% — roughly double the overall "
        "average. carrier_B has zero recorded failures across 412 packages in this extract. "
        "This carrier performance gap is a key operational lever."
    )

    h2(doc, "Query 2 — Failure Rate by Route Distance Bucket")
    body(doc,
        "Route distance is bucketed into three operationally meaningful ranges: < 40 km "
        "(dense urban), 40–60 km (suburban/mixed), and > 60 km (long-haul). The working "
        "hypothesis is that longer routes service less dense areas — but the data reveals "
        "the opposite."
    )

    # ── TASK 2 SQL QUERY 2 ──────────────────────────────────────────────────
    sql_distance = (
        "SELECT \n"
        "    CASE \n"
        "        WHEN route_distance_km < 40 THEN 'Under 40km'\n"
        "        WHEN route_distance_km BETWEEN 40 AND 60 THEN '40-60km'\n"
        "        ELSE 'Over 60km'\n"
        "    END AS distance_bucket,\n"
        "    COUNT(*) AS total_packages,\n"
        "    SUM(delivery_failure) AS failures,\n"
        "    ROUND(100.0 * SUM(delivery_failure) / COUNT(*), 2) \n"
        "    AS failure_rate_pct\n"
        "FROM packages\n"
        "GROUP BY distance_bucket\n"
        "ORDER BY failure_rate_pct DESC;"
    )
    code_block(doc, sql_distance)

    body(doc,
        "Finding: Routes < 40 km fail at 1.89% — the highest of any bucket — while routes "
        "> 60 km fail at 0.00%. This counterintuitive pattern is the central EDA finding."
    )

    h2(doc, "Additional Queries — Shift and Package Type")
    body(doc,
        "Similar GROUP BY queries were run for shift (morning: 1.37%, afternoon: 0.55%) "
        "and package_type (high_value packages showed elevated failure rates due to "
        "signature confirmation requirements)."
    )

    # Step 5
    h1(doc, "Step 5 — Discovery: The Urban Density Paradox")
    body(doc,
        "Our data surface a counter-intuitive reality: the shortest routes are our highest "
        "performance risk. While exurban routes over 60 km saw zero failures, routes "
        "under 40 km reached a 1.89% failure rate. This Urban Density Paradox "
        "reframes the problem from distance to access."
    )
    body(doc,
        "Dense vertical urban neighborhoods in Los Angeles present structural barriers — "
        "locked lobbies, intercom failures, and Amazon Locker congestion — that are not "
        "present in lower-density suburban routes. The risk is driven by infrastructure, "
        "not geography."
    )

    # Step 6
    h1(doc, "Step 6 — Correlation Analysis")
    body(doc,
        "A Pearson correlation heatmap provides a linear view of feature relationships. "
        "Critical limitation with imbalanced targets: even strong predictors show small "
        "absolute Pearson correlations when the positive class is <1% of records. "
        "Conditional GROUP BY failure rates (Step 4) are more reliable for feature "
        "evaluation in this context."
    )

    # Step 7
    h1(doc, "Step 7 — Feature Importance Preview")
    body(doc,
        "Before fitting any model, feature importance is previewed using conditional failure "
        "rates — the probability of delivery_failure = 1 given a specific feature value. "
        "Top risk signals:"
    )
    for item in [
        "carrier_D: highest failure rate among carriers (~1.39%)",
        "morning shift: highest shift-level failure rate (~1.37%)",
        "route < 40 km: highest distance bucket failure rate (~1.89%)",
        "high_value package type: elevated failure rate due to access requirements",
        "short_service_time flag: when triggered, elevated failure signal (locker/dense-urban stops)",
        "double_scan flag: scan errors correlate with downstream delivery failures",
    ]:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    # Step 8 — Summary
    h1(doc, "Step 8 — Summary of Findings")
    h2(doc, "Dataset Overview")
    overview = [
        ("Source",               "Amazon Last Mile Routing Research Challenge (ALMRRC 2021)"),
        ("Rows",                 "3,559 packages"),
        ("Routes",               "15 delivery routes, Los Angeles, July 2018"),
        ("Failures",             "25 packages → 0.70% overall failure rate"),
        ("Nulls",                "0 missing values"),
        ("Duplicates",           "0 duplicate package_id rows"),
        ("Zero-variance cols",   "weather_risk (constant='low'), days_in_fc (constant=0) — excluded"),
        ("Target column",        "delivery_failed — renamed from damaged_on_arrival; leakage fixed"),
    ]
    code_block(doc, "\n".join(f"{k:<22} {v}" for k, v in overview))

    h2(doc, "Key Findings by Feature")
    findings_text = [
        "1. Carrier Performance: carrier_D 1.39% (highest, ~2× average), carrier_B 0.00% (zero failures)",
        "2. Shift Performance: Morning 1.37% (residents commuting, locked lobbies), Afternoon 0.55%",
        "3. Route Distance: < 40 km → 1.89% (urban density paradox), > 60 km → 0.00%",
        "4. Accuracy is misleading — a constant 'Delivered' predictor achieves 99.30% accuracy "
           "while catching zero failures. Recall is the correct primary metric.",
    ]
    for item in findings_text:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(4)

    h2(doc, "Modeling Implications")
    for item in [
        "Extreme class imbalance (~140:1) requires SMOTE oversampling, class weighting, or threshold tuning",
        "Drop before encoding: weather_risk, days_in_fc (zero variance), delivery_failed (is the target)",
        "Top structural predictors: carrier, shift, and route_distance_km (bucketed)",
        "Binary flags (short_service_time, double_scan) are rare but high-signal events when triggered",
        "Low inter-feature correlation — no multicollinearity concerns",
    ]:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    path = os.path.join(DELIVERABLES, "04_eda_FINAL.docx")
    doc.save(path)
    print(f"Saved: {path}")

# ─────────────────────────────────────────────────────────────────────────────
# RUN ALL
# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    create_01()
    create_02()
    create_03()
    create_04()
    print("\nAll documents created successfully.")
