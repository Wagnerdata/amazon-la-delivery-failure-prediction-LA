from docx import Document
from docx.shared import Pt
import os

def create_resume_docx(md_path, docx_path):
    doc = Document()
    
    # Define basic style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for i, line in enumerate(lines):
        line = line.strip()
        if not line or line == "---":
            continue
        
        # Heading 1 (Name)
        if line.startswith("# "):
            h = doc.add_heading(line[2:], level=0)
            h.alignment = 1 # Center
        
        # Professional Title (usually line 2)
        elif i == 1 and "**" in line:
            p = doc.add_paragraph()
            p.alignment = 1 # Center
            line_clean = line.replace("**", "")
            run = p.add_run(line_clean)
            run.font.size = Pt(12)
            run.italic = True
        
        # Heading 2 (Sections)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        
        # Heading 3 (Jobs/Projects)
        elif line.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(line[4:])
            run.bold = True
            run.font.size = Pt(12)
        
        # Bullet points
        elif line.startswith("- "):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
        
        # Normal text or bold headers
        else:
            p = doc.add_paragraph()
            # Handle markdown bold **text**
            if "**" in line:
                parts = line.split("**")
                for j, part in enumerate(parts):
                    run = p.add_run(part)
                    if j % 2 != 0:
                        run.bold = True
            else:
                p.add_run(line)

    doc.save(docx_path)
    print(f"Resume saved to: {docx_path}")

if __name__ == "__main__":
    base_path = r"c:\Users\User\Correlation one Logistc Predict agent\delivery-failure-prediction"
    md_file = os.path.join(base_path, "reports", "Wagner_Campos_Professional_CV.md")
    # Using a final version name to avoid open-file locks
    docx_file = os.path.join(base_path, "Wagner_Campos_Professional_CV.docx")
    
    create_resume_docx(md_file, docx_file)
