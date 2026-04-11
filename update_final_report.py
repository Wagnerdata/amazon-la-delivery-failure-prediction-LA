"""
Task 3: Replace the GenAI Usage section in 07_final_report.docx.
"""
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

PATH = (r"C:\Users\User\Correlation one Logistc Predict agent"
        r"\delivery-failure-prediction\deliverables\07_final_report.docx")

doc = Document(PATH)
paras = doc.paragraphs

NEW_CONTENT = [
    ("The technical system — including the Random Forest classifier, CrewAI agent architecture, "
     "REST API, and data pipeline — was developed independently starting in October 2025, "
     "originally using synthetic delivery data. The system was designed from scratch to predict "
     "last-mile delivery failure before a package leaves the fulfillment center."),

    ("In April 2026, the system was adapted for the Correlation One Data Analytics program using "
     "real operational data from the Amazon Last Mile Routing Research Challenge "
     "(Los Angeles, 2018)."),

    ("The academic documents were first drafted in Jupyter notebooks to organize the analytical "
     "thinking and structure the arguments around the real data findings. The content was then "
     "formatted and refined in Word for formal delivery. Claude (Anthropic) assisted with writing, "
     "editing, and ensuring the documents met Correlation One requirements throughout this process."),

    ("All analysis and development was conducted in Jupyter notebooks, which are included in the "
     "project repository. The notebooks contain the full analytical pipeline \u2014 from data loading "
     "and SQL analysis to model training and evaluation \u2014 and can be executed independently to "
     "reproduce all results."),

    ("Claude (Anthropic) was used as a writing and editing tool for the academic deliverables. "
     "All analytical findings, technical decisions, and data interpretations are the author's own. "
     "CrewAI agents are part of the original system architecture \u2014 they generate per-package "
     "executive reports for dispatch supervisors, explaining model predictions in plain language."),

    ("This use of generative AI is consistent with Correlation One guidelines, which explicitly "
     "encourage transparent use of AI tools in the analytics workflow."),
]

# Find section boundaries
heading_idx = None
next_section_idx = None
for i, p in enumerate(paras):
    if "Generative AI Usage" in p.text and heading_idx is None:
        heading_idx = i
    if heading_idx is not None and i > heading_idx:
        if "4. Challenges" in p.text or "Challenges and Solutions" in p.text:
            next_section_idx = i
            break

first_body_idx = heading_idx + 1
print(f"Heading at:    {heading_idx}")
print(f"Body from:     {first_body_idx}")
print(f"Next section:  {next_section_idx}")
print(f"Removing paras: {list(range(first_body_idx, next_section_idx))}")

# Remove old body paragraphs (first_body_idx .. next_section_idx-1)
elements_to_remove = [paras[i]._element for i in range(first_body_idx, next_section_idx)]
ref_element = paras[next_section_idx]._element

for el in elements_to_remove:
    el.getparent().remove(el)

# Helper: build a normal paragraph XML element
def make_para(text):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'Normal')
    pPr.append(pStyle)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:after'), '120')
    pPr.append(spacing)
    p.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '22')   # 11pt
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p

# Insert in NORMAL order: each goes before ref_element, so they accumulate
# in correct sequence (0, 1, 2 ... each pushed right of previous)
for text in NEW_CONTENT:
    new_el = make_para(text)
    ref_element.addprevious(new_el)

doc.save(PATH)
print(f"\nSaved: {PATH}")

# Verify order
doc2 = Document(PATH)
print("\n--- GenAI section after update ---")
in_section = False
for p in doc2.paragraphs:
    if "Generative AI Usage" in p.text:
        in_section = True
    if in_section and ("4. Challenges" in p.text or "Challenges and Solutions" in p.text):
        break
    if in_section and p.text.strip():
        print(f"  {p.text[:110]}")
